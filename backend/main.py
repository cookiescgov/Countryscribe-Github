import os
import asyncio
import json
import glob
import re
import gc
import sys
import uuid
import time
import traceback
import subprocess
import threading
import smtplib
from email.message import EmailMessage
from datetime import datetime, timedelta
from typing import List, Optional
from io import BytesIO

from fastapi import FastAPI, UploadFile, File, Response, Form, HTTPException, Depends, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

import torch
import yt_dlp

# --- CUDA DIAGNOSTIC ---
print(f"--- SYSTEM DIAGNOSTIC ---", flush=True)
print(f"PyTorch version: {torch.__version__}", flush=True)
print(f"CUDA Available: {torch.cuda.is_available()}", flush=True)
if torch.cuda.is_available():
    print(f"CUDA Device: {torch.cuda.get_device_name(0)}", flush=True)
else:
    print(f"Reason for CPU: `torch.cuda.is_available()` is False. Check if NVIDIA Container Toolkit is passing the GPU.", flush=True)
print(f"-------------------------", flush=True)


# ─────────────────────────────────────────────
# Configuration
# ─────────────────────────────────────────────
ARCHIVE_ROOT = "archive"
os.makedirs(ARCHIVE_ROOT, exist_ok=True)

MAX_CONCURRENT_JOBS = int(os.getenv("MAX_CONCURRENT_JOBS", "2"))
SETTINGS_FILE = os.path.join(ARCHIVE_ROOT, "_settings.json")


def _default_smtp_settings() -> dict:
    """Seed from env vars so existing deployments keep working until they save in UI."""
    return {
        "smtp_host": os.getenv("SMTP_HOST", ""),
        "smtp_port": int(os.getenv("SMTP_PORT", "25") or 25),
        "smtp_user": os.getenv("SMTP_USER", ""),
        "smtp_pass": os.getenv("SMTP_PASS", ""),
        "smtp_from": os.getenv("SMTP_FROM", "County Scribe <noreply@localhost>"),
        "smtp_use_tls": (os.getenv("SMTP_USE_TLS", "").lower() in ("1", "true", "yes")),
        "app_base_url": os.getenv("APP_BASE_URL", "").rstrip("/"),
    }


def load_settings() -> dict:
    if os.path.exists(SETTINGS_FILE):
        try:
            with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
                saved = json.load(f)
            merged = _default_smtp_settings()
            merged.update({k: v for k, v in saved.items() if k in merged})
            return merged
        except Exception as e:
            print(f"Failed to read settings file: {e}", flush=True)
    return _default_smtp_settings()


def save_settings(new: dict) -> dict:
    current = load_settings()
    allowed = set(_default_smtp_settings().keys())
    for k, v in new.items():
        if k in allowed:
            current[k] = v
    # Types we care about
    try:
        current["smtp_port"] = int(current.get("smtp_port") or 25)
    except Exception:
        current["smtp_port"] = 25
    current["smtp_use_tls"] = bool(current.get("smtp_use_tls"))
    current["app_base_url"] = (current.get("app_base_url") or "").rstrip("/")
    with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
        json.dump(current, f, indent=2)
    try:
        os.chmod(SETTINGS_FILE, 0o600)
    except Exception:
        pass
    return current
JOB_RETENTION_SECONDS = 3600  # keep completed/failed jobs visible for 1h
HANG_THRESHOLD_SECONDS = 180  # UI may warn if heartbeat older than this (3 min)
JOBS_DIR = os.path.join(ARCHIVE_ROOT, "_jobs")
os.makedirs(JOBS_DIR, exist_ok=True)
PERSIST_THROTTLE_SECONDS = 2.0
MAX_RECENT_SEGMENTS_IN_STATE = 4000  # safety cap

# Semaphore: how many meetings may process concurrently
processing_sem = asyncio.Semaphore(MAX_CONCURRENT_JOBS)

# Job registry (in-memory). A restart clears it; any running job would be lost.
jobs: dict = {}

# Whisper model cache — keeping the model resident between jobs saves 30-60s per run.
# Thread lock because two concurrent jobs may race the first load.
_model_cache: dict = {}
_model_cache_lock = threading.Lock()

app = FastAPI()


@app.on_event("startup")
def _on_startup():
    _reconcile_persisted_jobs()
    _cleanup_old_jobs()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ─────────────────────────────────────────────
# Models
# ─────────────────────────────────────────────
class User(BaseModel):
    username: str


class TranscriptionSegment(BaseModel):
    start: str
    end: str
    text: str
    speaker: str = "Speaker"


# ─────────────────────────────────────────────
# Department DB (simple JSON file)
# ─────────────────────────────────────────────
USER_DB_FILE = os.path.join(ARCHIVE_ROOT, "users_db.json")


def load_users():
    if not os.path.exists(USER_DB_FILE):
        defaults = {
            "Planning Commission": {"username": "Planning Commission"},
            "Council/Commissioners": {"username": "Council/Commissioners"},
            "Drainage Board": {"username": "Drainage Board"},
            "Park Board": {"username": "Park Board"},
            "Health Board": {"username": "Health Board"},
        }
        try:
            with open(USER_DB_FILE, "w", encoding="utf-8") as f:
                json.dump(defaults, f, indent=2)
            os.chmod(USER_DB_FILE, 0o666)
        except Exception:
            pass
        return defaults

    try:
        with open(USER_DB_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def save_user_to_db(username: str):
    users = load_users()
    if username not in users:
        users[username] = {"username": username}
        with open(USER_DB_FILE, "w", encoding="utf-8") as f:
            json.dump(users, f, indent=2)
        try:
            os.chmod(USER_DB_FILE, 0o666)
        except Exception:
            pass


def get_user_email(username: str) -> Optional[str]:
    users = load_users()
    u = users.get(username) or {}
    return u.get("email") or None


def set_user_email(username: str, email: Optional[str]):
    users = load_users()
    if username not in users:
        users[username] = {"username": username}
    if email:
        users[username]["email"] = email.strip()
    else:
        users[username].pop("email", None)
    with open(USER_DB_FILE, "w", encoding="utf-8") as f:
        json.dump(users, f, indent=2)


# ─────────────────────────────────────────────
# Email helper
# ─────────────────────────────────────────────
def send_email(to_addr: str, subject: str, body: str, settings: Optional[dict] = None) -> tuple[bool, str]:
    s_cfg = settings if settings is not None else load_settings()
    host = (s_cfg.get("smtp_host") or "").strip()
    if not host:
        msg = "SMTP not configured; skipping email."
        print(msg, flush=True)
        return False, msg
    try:
        em = EmailMessage()
        em["From"] = s_cfg.get("smtp_from") or "County Scribe <noreply@localhost>"
        em["To"] = to_addr
        em["Subject"] = subject
        em.set_content(body)
        port = int(s_cfg.get("smtp_port") or 25)
        with smtplib.SMTP(host, port, timeout=30) as server:
            if s_cfg.get("smtp_use_tls"):
                server.starttls()
            user = (s_cfg.get("smtp_user") or "").strip()
            pw = s_cfg.get("smtp_pass") or ""
            if user and pw:
                server.login(user, pw)
            server.send_message(em)
        ok = f"Email sent to {to_addr}: {subject}"
        print(ok, flush=True)
        return True, ok
    except Exception as e:
        err = f"Email send failed: {type(e).__name__}: {e}"
        print(f"{err} (to={to_addr})", flush=True)
        return False, err


def _notify_job_result(job: dict):
    """Best-effort email when a job reaches a terminal state."""
    email = job.get("notify_email")
    if not email:
        return
    state = job.get("state")
    fname = job.get("filename") or "your meeting"
    s_cfg = load_settings()
    app_base = (s_cfg.get("app_base_url") or "").rstrip("/")
    app_link = app_base if app_base else "(set App Base URL in Settings to include a link)"

    # Notify on complete or failed (but not on user-initiated cancel)
    if state == "complete":
        subject = f"[County Scribe] Transcription ready: {fname}"
        body = (
            f"Your transcription is complete.\n\n"
            f"File: {fname}\n"
            f"Archive: {job.get('archive_filename')}\n"
            f"Department: {job.get('username')}\n"
            f"Segments: {job.get('segments_done')}\n"
            f"Elapsed: {int((job.get('ended_at') or 0) - (job.get('started_at') or 0))}s\n\n"
            f"Open County Scribe: {app_link}\n"
        )
    elif state == "failed":
        subject = f"[County Scribe] Transcription FAILED: {fname}"
        body = (
            f"Your transcription did not complete.\n\n"
            f"File: {fname}\n"
            f"Error code: {job.get('error_code')}\n"
            f"Error: {job.get('error')}\n"
            f"Department: {job.get('username')}\n\n"
            f"Open County Scribe: {app_link}\n"
        )
    else:
        return
    send_email(email, subject, body)


# ─────────────────────────────────────────────
# Identity helpers (Auth-Free Department Tagging)
# ─────────────────────────────────────────────
async def get_current_user(x_department: str = Header(default="Unknown_Department")) -> User:
    return User(username=x_department)


# ─────────────────────────────────────────────
# Archive helpers
# ─────────────────────────────────────────────
def get_user_archive_dir(username: str) -> str:
    safe_name = re.sub(r"[^a-zA-Z0-9]", "_", username)
    user_dir = os.path.join(ARCHIVE_ROOT, safe_name)
    if not os.path.exists(user_dir):
        os.makedirs(user_dir, exist_ok=True)
        try:
            os.chmod(user_dir, 0o777)
        except Exception:
            pass
    return user_dir


def format_timestamp(seconds: float) -> str:
    if seconds is None:
        return "00:00:00.000"
    td = timedelta(seconds=seconds)
    hours, remainder = divmod(int(td.total_seconds()), 3600)
    minutes, sec = divmod(remainder, 60)
    return f"{hours:02}:{minutes:02}:{sec:02}.{td.microseconds // 1000:03}"


# ─────────────────────────────────────────────
# Job helpers
# ─────────────────────────────────────────────
def _now() -> float:
    return time.time()


def _new_job(username: str, original_filename: str, source: str,
             notify_email: Optional[str] = None) -> str:
    job_id = uuid.uuid4().hex[:12]
    jobs[job_id] = {
        "id": job_id,
        "username": username,
        "filename": original_filename,
        "source": source,          # "file" | "youtube"
        "state": "queued",         # queued|downloading|normalizing|loading_model|transcribing|archiving|complete|failed|cancelled
        "phase": "Queued",         # human-readable
        "progress": 0.0,           # 0.0 – 1.0
        "message": "",             # latest detail line
        "error": None,
        "error_code": None,
        "error_detail": None,
        "started_at": _now(),
        "ended_at": None,
        "last_heartbeat": _now(),
        "audio_duration": None,
        "segments_done": 0,
        "archive_filename": None,
        "cancel_requested": False,
        "queue_position": None,
        "notify_email": (notify_email or None),
        "segments": [],            # formatted partial transcript (in-memory only)
        "_last_persist": 0.0,
    }
    _persist_job(job_id, force=True)
    return job_id


# --- Non-persisted fields (not written to disk) ---
_NON_PERSISTED = {"segments", "_last_persist", "cancel_requested"}


def _persist_job(job_id: str, force: bool = False):
    j = jobs.get(job_id)
    if not j:
        return
    now = _now()
    if not force and (now - j.get("_last_persist", 0.0)) < PERSIST_THROTTLE_SECONDS:
        return
    try:
        snapshot = {k: v for k, v in j.items() if k not in _NON_PERSISTED}
        path = os.path.join(JOBS_DIR, f"{job_id}.json")
        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(snapshot, f, ensure_ascii=False)
        os.replace(tmp, path)
        j["_last_persist"] = now
    except Exception as e:
        print(f"Persist failed for {job_id}: {e}", flush=True)


def _update(job_id: str, **kwargs):
    j = jobs.get(job_id)
    if not j:
        return
    state_changed = "state" in kwargs and kwargs["state"] != j.get("state")
    j.update(kwargs)
    j["last_heartbeat"] = _now()
    _persist_job(job_id, force=state_changed)


def _heartbeat(job_id: str):
    j = jobs.get(job_id)
    if j:
        j["last_heartbeat"] = _now()


def _reconcile_persisted_jobs():
    """On startup, load persisted jobs. Any non-terminal state = server crashed mid-job."""
    if not os.path.isdir(JOBS_DIR):
        return
    for fname in os.listdir(JOBS_DIR):
        if not fname.endswith(".json"):
            continue
        path = os.path.join(JOBS_DIR, fname)
        try:
            with open(path, "r", encoding="utf-8") as f:
                j = json.load(f)
        except Exception:
            continue
        # Rehydrate, filling missing fields
        j.setdefault("segments", [])
        j.setdefault("_last_persist", 0.0)
        j.setdefault("cancel_requested", False)
        jid = j.get("id")
        if not jid:
            continue
        if j.get("state") not in ("complete", "failed", "cancelled"):
            j["state"] = "failed"
            j["phase"] = "Failed (server restart)"
            j["error"] = "Server restarted while this job was running."
            j["error_code"] = "SERVER_RESTART"
            j["error_detail"] = None
            j["ended_at"] = _now()
            try:
                snap = {k: v for k, v in j.items() if k not in _NON_PERSISTED}
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(snap, f, ensure_ascii=False)
            except Exception:
                pass
        jobs[jid] = j
    print(f"Reconciled {len(jobs)} persisted jobs.", flush=True)


def _classify_error(exc: BaseException) -> str:
    msg = f"{type(exc).__name__}: {exc}".lower()
    if "cuda out of memory" in msg or "cublas_status_alloc_failed" in msg:
        return "CUDA_OOM"
    if "ffmpeg" in msg:
        return "FFMPEG_FAIL"
    if "yt" in msg or "youtube" in msg or "http error" in msg:
        return "YT_DOWNLOAD_FAIL"
    if "model" in msg and ("load" in msg or "download" in msg):
        return "MODEL_LOAD_FAIL"
    if isinstance(exc, asyncio.CancelledError):
        return "CANCELLED"
    if isinstance(exc, FileNotFoundError):
        return "FILE_NOT_FOUND"
    return "UNKNOWN"


def _cleanup_old_jobs():
    now = _now()
    for jid in list(jobs.keys()):
        j = jobs[jid]
        if j["state"] in ("complete", "failed", "cancelled"):
            ended = j.get("ended_at") or j["started_at"]
            if now - ended > JOB_RETENTION_SECONDS:
                del jobs[jid]
                try:
                    os.remove(os.path.join(JOBS_DIR, f"{jid}.json"))
                except FileNotFoundError:
                    pass
                except Exception:
                    pass


def _recalc_queue_positions():
    queued = [j for j in jobs.values() if j["state"] == "queued"]
    queued.sort(key=lambda j: j["started_at"])
    for i, j in enumerate(queued):
        j["queue_position"] = i + 1


# ─────────────────────────────────────────────
# Health
# ─────────────────────────────────────────────
@app.get("/api/health")
def health():
    return {
        "message": "County Scribe API is running.",
        "max_concurrent_jobs": MAX_CONCURRENT_JOBS,
        "active_jobs": sum(1 for j in jobs.values() if j["state"] not in ("complete", "failed", "cancelled")),
    }


# ─────────────────────────────────────────────
# Department User Management
# ─────────────────────────────────────────────
@app.get("/api/users")
def get_all_users():
    users = load_users()
    return list(users.keys())


@app.post("/api/register")
async def register_department(username: str = Form(...)):
    save_user_to_db(username.strip())
    return {"message": "Department created", "username": username.strip()}


@app.delete("/api/users/{username}")
async def delete_department(username: str):
    users = load_users()
    if username in users:
        del users[username]
        with open(USER_DB_FILE, "w", encoding="utf-8") as f:
            json.dump(users, f, indent=2)
        return {"message": "Department deleted"}
    raise HTTPException(status_code=404, detail="Department not found")


# ─────────────────────────────────────────────
# Archive endpoints (per-user)
# ─────────────────────────────────────────────
@app.get("/api/archives")
def list_archives(current_user: User = Depends(get_current_user)):
    user_dir = get_user_archive_dir(current_user.username)
    files = glob.glob(os.path.join(user_dir, "*.json"))

    archives = []
    for fpath in files:
        stats = os.stat(fpath)
        filename = os.path.basename(fpath)

        meeting_date = "Unknown"
        match = re.match(r"(\d{4}-\d{2}-\d{2})", filename)
        if match:
            meeting_date = match.group(1)

        archives.append(
            {
                "filename": filename,
                "created": datetime.fromtimestamp(stats.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
                "meeting_date": meeting_date,
                "size_kb": round(stats.st_size / 1024, 2),
            }
        )

    return sorted(archives, key=lambda x: x["filename"], reverse=True)


@app.get("/api/archives/{filename}")
def get_archive(filename: str, current_user: User = Depends(get_current_user)):
    user_dir = get_user_archive_dir(current_user.username)
    safe_path = os.path.join(user_dir, os.path.basename(filename))

    if not os.path.exists(safe_path):
        raise HTTPException(status_code=404, detail="Archive not found")

    with open(safe_path, "r", encoding="utf-8") as f:
        return json.load(f)


@app.delete("/api/archives/{filename}")
def delete_archive(filename: str, current_user: User = Depends(get_current_user)):
    user_dir = get_user_archive_dir(current_user.username)
    safe_path = os.path.join(user_dir, os.path.basename(filename))

    if os.path.exists(safe_path):
        os.remove(safe_path)
        return {"status": "deleted"}

    raise HTTPException(status_code=404, detail="File not found")


class RenamePayload(BaseModel):
    new_name: str


@app.patch("/api/archives/{filename}")
def rename_archive(filename: str, payload: RenamePayload,
                   current_user: User = Depends(get_current_user)):
    user_dir = get_user_archive_dir(current_user.username)
    src = os.path.join(user_dir, os.path.basename(filename))
    if not os.path.exists(src):
        raise HTTPException(status_code=404, detail="Archive not found")

    # Sanitize new name; force .json; strip path bits
    base = os.path.basename(payload.new_name.strip())
    if not base:
        raise HTTPException(status_code=400, detail="New name cannot be empty.")
    base = re.sub(r"[^A-Za-z0-9._ -]", "_", base)
    if not base.lower().endswith(".json"):
        base = base + ".json"
    dest = os.path.join(user_dir, base)
    if os.path.exists(dest) and os.path.abspath(dest) != os.path.abspath(src):
        raise HTTPException(status_code=409, detail="A file with that name already exists.")
    os.rename(src, dest)
    return {"status": "renamed", "old": os.path.basename(src), "new": base}


# ─────────────────────────────────────────────
# Per-department email (for notifications)
# ─────────────────────────────────────────────
class EmailPayload(BaseModel):
    email: Optional[str] = None


@app.get("/api/me/email")
def get_my_email(current_user: User = Depends(get_current_user)):
    return {"username": current_user.username, "email": get_user_email(current_user.username)}


@app.put("/api/me/email")
def put_my_email(payload: EmailPayload, current_user: User = Depends(get_current_user)):
    email = (payload.email or "").strip() or None
    if email and not re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email):
        raise HTTPException(status_code=400, detail="Invalid email address.")
    set_user_email(current_user.username, email)
    return {"username": current_user.username, "email": email}


# ─────────────────────────────────────────────
# Server settings (SMTP / app base URL)
# ─────────────────────────────────────────────
class SmtpSettingsPayload(BaseModel):
    smtp_host: Optional[str] = None
    smtp_port: Optional[int] = None
    smtp_user: Optional[str] = None
    smtp_pass: Optional[str] = None
    smtp_from: Optional[str] = None
    smtp_use_tls: Optional[bool] = None
    app_base_url: Optional[str] = None


class SmtpTestPayload(BaseModel):
    to: str


def _mask_settings_for_display(s: dict) -> dict:
    """Never expose the password over the API; flag whether one is set."""
    out = dict(s)
    pw = out.pop("smtp_pass", "") or ""
    out["smtp_pass_set"] = bool(pw)
    return out


@app.get("/api/settings/smtp")
def get_smtp_settings(current_user: User = Depends(get_current_user)):
    return _mask_settings_for_display(load_settings())


@app.put("/api/settings/smtp")
def put_smtp_settings(payload: SmtpSettingsPayload, current_user: User = Depends(get_current_user)):
    data = {k: v for k, v in payload.dict().items() if v is not None}
    # Empty string for smtp_pass explicitly clears it; None means leave alone.
    saved = save_settings(data)
    return _mask_settings_for_display(saved)


@app.post("/api/settings/smtp/test")
def test_smtp_settings(payload: SmtpTestPayload, current_user: User = Depends(get_current_user)):
    to = (payload.to or "").strip()
    if not re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", to):
        raise HTTPException(status_code=400, detail="Provide a valid test recipient email.")
    cfg = load_settings()
    if not (cfg.get("smtp_host") or "").strip():
        raise HTTPException(status_code=400, detail="SMTP host is not set. Save settings first.")
    ok, msg = send_email(
        to,
        "[County Scribe] SMTP test message",
        "This is a test email from County Scribe. If you got this, your mail settings are working.\n",
        settings=cfg,
    )
    if not ok:
        raise HTTPException(status_code=502, detail=msg)
    return {"ok": True, "message": msg}


@app.post("/api/archives/prune")
def prune_archives(days: int = 180, current_user: User = Depends(get_current_user)):
    cutoff = (datetime.now() - timedelta(days=days)).timestamp()
    user_dir = get_user_archive_dir(current_user.username)

    deleted_count = 0
    files = glob.glob(os.path.join(user_dir, "*.json"))
    for fpath in files:
        if os.stat(fpath).st_mtime < cutoff:
            os.remove(fpath)
            deleted_count += 1

    return {"status": "success", "deleted": deleted_count, "retention_days": days}


# ─────────────────────────────────────────────
# Audio helpers (sync; run in worker thread)
# ─────────────────────────────────────────────
def _fmt_bytes(n):
    if not n:
        return "?"
    for unit in ("B", "KB", "MB", "GB"):
        if n < 1024:
            return f"{n:.1f} {unit}"
        n /= 1024
    return f"{n:.1f} TB"


def download_youtube_audio(url: str, job_id: Optional[str] = None) -> tuple[str, str]:
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    output_template = f"temp_yt_{timestamp}"

    def progress_hook(d):
        if not job_id:
            return
        status = d.get("status")
        if status == "downloading":
            total = d.get("total_bytes") or d.get("total_bytes_estimate") or 0
            got = d.get("downloaded_bytes", 0)
            pct = (got / total) if total else 0.0
            # Map YT download into 2%..8% of overall progress
            overall = 0.02 + 0.06 * pct
            msg = f"Downloading: {_fmt_bytes(got)}"
            if total:
                msg += f" / {_fmt_bytes(total)} ({int(pct*100)}%)"
            spd = d.get("speed")
            if spd:
                msg += f" — {_fmt_bytes(spd)}/s"
            _update(job_id, progress=overall, message=msg)
        elif status == "finished":
            _update(job_id, progress=0.08, message="Download complete, extracting audio…")

    ydl_opts = {
        "format": "bestaudio/best",
        "outtmpl": output_template,
        "postprocessors": [{"key": "FFmpegExtractAudio", "preferredcodec": "m4a"}],
        "quiet": True,
        "no_warnings": True,
        "source_address": "0.0.0.0",
        "socket_timeout": 60,
        "retries": 20,
        "fragment_retries": 20,
        "concurrent_fragment_downloads": 4,
        "progress_hooks": [progress_hook],
    }

    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=False)
        title = info.get("title", "YouTube_Video")
        print(f"Starting download: {title}", flush=True)
        if job_id:
            _update(job_id, message=f"Starting: {title}")
        ydl.download([url])
        final_filename = f"{output_template}.m4a"
        return final_filename, title


def normalize_audio(input_path: str) -> str:
    output_path = input_path + ".normalized.wav"
    try:
        print(f"Normalizing audio: {input_path} -> {output_path}", flush=True)
        command = [
            "ffmpeg", "-y", "-i", input_path,
            "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le",
            output_path
        ]
        result = subprocess.run(command, capture_output=True, text=True)
        if result.returncode != 0:
            print(f"FFmpeg Error: {result.stderr}", flush=True)
            return input_path
        return output_path
    except Exception as e:
        print(f"Normalization exception: {e}", flush=True)
        return input_path


# ─────────────────────────────────────────────
# Transcription worker (runs in a thread)
# ─────────────────────────────────────────────
class JobCancelled(Exception):
    pass


def _check_cancel(job_id: str):
    j = jobs.get(job_id)
    if j and j.get("cancel_requested"):
        raise JobCancelled()


def _run_transcription_sync(job_id: str, temp_audio_path: str, model_size: str,
                            original_filename: str, meeting_date: Optional[str], username: str,
                            youtube_url: Optional[str]) -> str:
    """
    The long-running transcription pipeline. Runs in a worker thread.
    Updates the job dict in-place. Returns the archive filename on success.
    """
    try:
        # 1) If YouTube, download first
        if youtube_url:
            _update(job_id, state="downloading", phase="Downloading YouTube audio", progress=0.02)
            temp_audio_path, video_title = download_youtube_audio(youtube_url, job_id=job_id)
            safe_title = "".join([c for c in video_title if c.isalnum() or c in (" ", ".", "_")]).replace(" ", "_")
            original_filename = f"{safe_title}.m4a"
            _update(job_id, filename=original_filename)

        _check_cancel(job_id)

        # 2) Normalize
        _update(job_id, state="normalizing", phase="Preparing audio (ffmpeg)", progress=0.08)
        processed_audio_path = normalize_audio(temp_audio_path)

        _check_cancel(job_id)

        # 3) Load (or reuse) model
        from faster_whisper import WhisperModel
        device = "cuda" if torch.cuda.is_available() else "cpu"
        if device == "cuda":
            torch.cuda.empty_cache()
        compute_type = "float16" if (model_size == "large-v3" and device == "cuda") else "int8"

        cache_key = f"{model_size}|{device}|{compute_type}"
        with _model_cache_lock:
            model = _model_cache.get(cache_key)
            if model is None:
                _update(job_id, state="loading_model",
                        phase=f"Loading {model_size} ({device}/{compute_type})", progress=0.12,
                        message="First-time model load…")
                model = WhisperModel(model_size, device=device, compute_type=compute_type)
                _model_cache[cache_key] = model
            else:
                _update(job_id, state="loading_model",
                        phase=f"Model ready ({model_size})", progress=0.14,
                        message="Using cached model — skipping load")

        _check_cancel(job_id)

        # 4) Transcribe — stream segments for real progress.
        # vad_filter=True skips silence/non-speech: typically 20-40% faster on meetings
        # with pauses, with no loss of transcription quality.
        _update(job_id, state="transcribing", phase="Transcribing audio", progress=0.15,
                message="Warming up decoder")
        segments, info = model.transcribe(
            processed_audio_path,
            beam_size=5,
            vad_filter=True,
            vad_parameters=dict(min_silence_duration_ms=500),
        )
        duration = getattr(info, "duration", None) or 0.0
        _update(job_id, audio_duration=duration)

        result_segments = []
        for segment in segments:
            _check_cancel(job_id)
            result_segments.append({
                "start": segment.start,
                "end": segment.end,
                "text": segment.text,
                "speaker": "Speaker",
            })
            n = len(result_segments)
            if duration > 0:
                pct = 0.15 + 0.80 * min(1.0, (segment.end or 0.0) / duration)
            else:
                pct = min(0.95, 0.15 + n * 0.001)

            # Append formatted segment for live partial-transcript streaming
            j = jobs.get(job_id)
            if j is not None and len(j["segments"]) < MAX_RECENT_SEGMENTS_IN_STATE:
                j["segments"].append({
                    "start": format_timestamp(segment.start or 0),
                    "end": format_timestamp(segment.end or 0),
                    "text": (segment.text or "").strip(),
                    "speaker": "Speaker",
                })

            _update(job_id, progress=pct, segments_done=n,
                    message=f"Transcribed {n} segments ({format_timestamp(segment.end or 0)})")

        _update(job_id, progress=0.96, message=f"Transcription complete — {len(result_segments)} segments")

        # Model stays cached for the next job. Just trim transient working-set memory.
        gc.collect()

        # 5) Format and archive
        _update(job_id, state="archiving", phase="Saving transcript", progress=0.98)
        formatted_result = [
            {
                "start": format_timestamp(seg.get("start", 0)),
                "end": format_timestamp(seg.get("end", 0)),
                "text": (seg.get("text", "") or "").strip(),
                "speaker": "Speaker",
            }
            for seg in result_segments
        ]

        date_prefix = meeting_date if meeting_date else datetime.now().strftime("%Y-%m-%d")
        time_suffix = datetime.now().strftime("%H-%M-%S")
        safe_filename = "".join([c for c in original_filename if c.isalnum() or c in (" ", ".", "_")]).replace(" ", "_")

        user_dir = get_user_archive_dir(username)
        archive_name = f"{date_prefix}_{time_suffix}_{safe_filename}.json"
        archive_path = os.path.join(user_dir, archive_name)

        with open(archive_path, "w", encoding="utf-8") as f:
            json.dump(formatted_result, f, indent=2, ensure_ascii=False)
        print(f"Successfully archived to: {archive_path}", flush=True)

        _update(job_id, state="complete", phase="Complete", progress=1.0,
                archive_filename=archive_name, ended_at=_now(),
                message=f"Saved {archive_name}")
        return archive_name

    finally:
        for p in [temp_audio_path, (temp_audio_path + ".normalized.wav" if temp_audio_path else None)]:
            if p and os.path.exists(p):
                try:
                    os.remove(p)
                    print(f"Cleaned up temp file: {p}", flush=True)
                except Exception:
                    pass


async def _run_job(job_id: str, temp_audio_path: Optional[str], model_size: str,
                   original_filename: str, meeting_date: Optional[str], username: str,
                   youtube_url: Optional[str]):
    _recalc_queue_positions()
    try:
        async with processing_sem:
            j = jobs.get(job_id)
            if j and j.get("cancel_requested"):
                _update(job_id, state="cancelled", phase="Cancelled before start",
                        ended_at=_now())
                _recalc_queue_positions()
                return
            _update(job_id, queue_position=0, phase="Starting", state="normalizing",
                    progress=0.01)
            _recalc_queue_positions()
            await asyncio.to_thread(
                _run_transcription_sync, job_id, temp_audio_path, model_size,
                original_filename, meeting_date, username, youtube_url,
            )
    except JobCancelled:
        _update(job_id, state="cancelled", phase="Cancelled by user",
                ended_at=_now(), message="Stopped by user")
    except Exception as e:
        code = _classify_error(e)
        tb = traceback.format_exc()
        print(f"JOB {job_id} FAILED [{code}]: {e}\n{tb}", flush=True)
        _update(job_id, state="failed", phase=f"Failed ({code})",
                error=str(e), error_code=code,
                error_detail=tb[-4000:], ended_at=_now())
    finally:
        # Best-effort notification on terminal state
        try:
            j = jobs.get(job_id)
            if j:
                _notify_job_result(j)
        except Exception as _e:
            print(f"Notify failed for {job_id}: {_e}", flush=True)
        _cleanup_old_jobs()
        _recalc_queue_positions()


# ─────────────────────────────────────────────
# Transcribe endpoint — returns a job_id immediately
# ─────────────────────────────────────────────
@app.post("/api/transcribe")
async def transcribe_audio(
    file: UploadFile = File(None),
    youtube_url: str = Form(None),
    model_size: str = Form("large-v3"),
    diarize: bool = Form(False),
    meeting_date: str = Form(None),
    notify_email: str = Form(None),
    current_user: User = Depends(get_current_user),
):
    if not file and not youtube_url:
        raise HTTPException(status_code=400, detail="Please provide either a file or a YouTube URL.")

    # Save upload to disk synchronously (file handle dies with the request)
    temp_audio_path = None
    original_filename = "YouTube_Video"
    source = "youtube" if youtube_url else "file"

    if file:
        original_filename = file.filename or "uploaded_audio"
        temp_audio_path = f"temp_{uuid.uuid4().hex[:8]}_{original_filename}"
        print(f"Saving upload to {temp_audio_path}...", flush=True)
        content = await file.read()
        with open(temp_audio_path, "wb") as f:
            f.write(content)
        print(f"Upload saved ({len(content)} bytes).", flush=True)

    # Resolve notification email: explicit form value wins, else department default
    email = (notify_email or "").strip() or None
    if email and not re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email):
        raise HTTPException(status_code=400, detail="Invalid notify_email.")
    if email:
        # Persist as the department's email for future confirmations
        set_user_email(current_user.username, email)

    job_id = _new_job(current_user.username, original_filename, source, notify_email=email)
    _recalc_queue_positions()

    asyncio.create_task(_run_job(
        job_id, temp_audio_path, model_size, original_filename, meeting_date,
        current_user.username, youtube_url,
    ))

    return {"job_id": job_id, "state": jobs[job_id]["state"]}


# ─────────────────────────────────────────────
# Job status / cancel
# ─────────────────────────────────────────────
def _job_view(j: dict) -> dict:
    now = _now()
    elapsed = now - j["started_at"]
    eta = None
    if j["state"] == "transcribing" and j["progress"] > 0.2 and j["progress"] < 1.0:
        # Extrapolate from current progress
        total = elapsed / max(j["progress"], 0.01)
        eta = max(0, int(total - elapsed))
    stale = (now - j["last_heartbeat"]) if j["state"] not in ("complete", "failed", "cancelled") else 0
    return {
        "id": j["id"],
        "state": j["state"],
        "phase": j["phase"],
        "progress": round(j["progress"], 4),
        "message": j["message"],
        "filename": j["filename"],
        "source": j["source"],
        "segments_done": j["segments_done"],
        "audio_duration": j["audio_duration"],
        "queue_position": j["queue_position"],
        "error": j["error"],
        "error_code": j["error_code"],
        "error_detail": j["error_detail"],
        "archive_filename": j["archive_filename"],
        "started_at": j["started_at"],
        "ended_at": j["ended_at"],
        "elapsed_seconds": int(elapsed),
        "eta_seconds": eta,
        "seconds_since_heartbeat": int(stale),
        "possibly_hung": stale > HANG_THRESHOLD_SECONDS,
    }


@app.get("/api/jobs")
def list_jobs(current_user: User = Depends(get_current_user)):
    out = [_job_view(j) for j in jobs.values() if j["username"] == current_user.username]
    out.sort(key=lambda x: x["started_at"], reverse=True)
    return out


@app.get("/api/jobs/{job_id}")
def get_job(job_id: str, current_user: User = Depends(get_current_user)):
    j = jobs.get(job_id)
    if not j:
        raise HTTPException(status_code=404, detail="Job not found (may have been cleaned up after 1h).")
    return _job_view(j)


@app.get("/api/jobs/{job_id}/segments")
def get_job_segments(job_id: str, after: int = 0, current_user: User = Depends(get_current_user)):
    """
    Returns transcribed segments with index > `after`.
    Frontend keeps its own count and passes the highest seen index, so only new ones come over the wire.
    """
    j = jobs.get(job_id)
    if not j:
        raise HTTPException(status_code=404, detail="Job not found.")
    segs = j.get("segments", [])
    total = len(segs)
    if after < 0:
        after = 0
    new_segs = segs[after:] if after < total else []
    return {"total": total, "after": after, "segments": new_segs}


@app.delete("/api/jobs/{job_id}")
def cancel_job(job_id: str, current_user: User = Depends(get_current_user)):
    j = jobs.get(job_id)
    if not j:
        raise HTTPException(status_code=404, detail="Job not found.")
    if j["state"] in ("complete", "failed", "cancelled"):
        return _job_view(j)
    j["cancel_requested"] = True
    j["message"] = "Cancellation requested…"
    _heartbeat(job_id)
    return _job_view(j)


# ─────────────────────────────────────────────
# Export endpoints
# ─────────────────────────────────────────────
@app.post("/api/download-pdf")
async def download_pdf(transcription_data: List[TranscriptionSegment], clean: bool = False):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [Paragraph("Meeting Minutes" if clean else "Transcription", styles["h1"]), Spacer(1, 20)]

    if clean:
        block = []
        for i, segment in enumerate(transcription_data):
            block.append(segment.text.strip())
            if (i + 1) % 5 == 0:
                story.append(Paragraph(" ".join(block), styles["Normal"]))
                story.append(Spacer(1, 10))
                block = []
        if block:
            story.append(Paragraph(" ".join(block), styles["Normal"]))
    else:
        for segment in transcription_data:
            story.append(Paragraph(f"<b>[{segment.start}] {segment.speaker}:</b> {segment.text}", styles["Normal"]))
            story.append(Spacer(1, 10))

    doc.build(story)
    buffer.seek(0)
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="minutes.pdf"'},
    )


@app.post("/api/download-docx")
async def download_docx(transcription_data: List[TranscriptionSegment], clean: bool = False):
    document = Document()
    document.add_heading("Meeting Minutes" if clean else "Transcription", 1)

    if clean:
        block = []
        for i, segment in enumerate(transcription_data):
            block.append(segment.text.strip())
            if (i + 1) % 5 == 0:
                document.add_paragraph(" ".join(block))
                block = []
        if block:
            document.add_paragraph(" ".join(block))
    else:
        for segment in transcription_data:
            document.add_paragraph(f"[{segment.start}] {segment.speaker}: {segment.text}")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="minutes.docx"'},
    )


# ─────────────────────────────────────────────
# Static frontend hosting (React build)
# ─────────────────────────────────────────────
if os.path.exists("static"):
    app.mount("/static", StaticFiles(directory="static/static"), name="static")

    @app.get("/{full_path:path}")
    async def serve_react(full_path: str):
        if full_path.startswith("api"):
            return {"error": "API route not found"}
        return FileResponse("static/index.html")
else:
    print("Warning: 'static' directory not found. Frontend will not be served.")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
