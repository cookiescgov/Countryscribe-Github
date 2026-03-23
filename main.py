import os
import asyncio
import json
import time
import glob
from datetime import datetime, timedelta
from typing import List, Optional
from io import BytesIO

from fastapi import FastAPI, UploadFile, File, Response, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

import torch
import gc
import yt_dlp
import re

# --- Configuration ---
# MOCK_MODE: Set to True for testing without GPU/Models
MOCK_MODE = os.getenv("MOCK_MODE", "False").lower() == "true"
ARCHIVE_DIR = "archive"
os.makedirs(ARCHIVE_DIR, exist_ok=True)

# Global lock to prevent concurrent GPU processing
processing_lock = asyncio.Lock()

# --- FastAPI App Initialization ---
app = FastAPI()

# --- CORS Configuration ---
origins = ["*"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class TranscriptionSegment(BaseModel):
    start: str
    end: str
    text: str

# --- Helper Function ---
def format_timestamp(seconds: float) -> str:
    """Converts a float timestamp in seconds to a formatted string HH:MM:SS.ms"""
    if seconds is None:
        return "00:00:00.000"
    td = timedelta(seconds=seconds)
    total_seconds = int(td.total_seconds())
    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    milliseconds = td.microseconds // 1000
    return f"{hours:02}:{minutes:02}:{seconds:02}.{milliseconds:03}"

# --- API Endpoints ---
@app.get("/api/health")
def read_root():
    return {"message": "County Scribe API is running (LXC Build - No Diarization)."}

# --- ARCHIVE ENDPOINTS ---

def cleanup_old_archives(days: int = 180):
    """Internal function to delete files older than X days."""
    try:
        cutoff = time.time() - (days * 86400)
        files = glob.glob(os.path.join(ARCHIVE_DIR, "*.json"))
        for f in files:
            if os.stat(f).st_mtime < cutoff:
                os.remove(f)
                print(f"Auto-Pruned expired archive: {f}")
    except Exception as e:
        print(f"Error during auto-prune: {e}")

@app.on_event("startup")
async def startup_event():
    """Run cleanup on server launch."""
    cleanup_old_archives()

@app.get("/api/archives")
def list_archives():
    """List all archived transcripts sorted by date (newest first). Auto-prunes old files."""
    cleanup_old_archives()
    
    files = glob.glob(os.path.join(ARCHIVE_DIR, "*.json"))
    archives = []
    for f in files:
        stats = os.stat(f)
        filename = os.path.basename(f)
        
        meeting_date = "Unknown"
        match = re.match(r"(\d{4}-\d{2}-\d{2})", filename)
        if match:
            meeting_date = match.group(1)

        archives.append({
            "filename": filename,
            "created": datetime.fromtimestamp(stats.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
            "meeting_date": meeting_date,
            "size_kb": round(stats.st_size / 1024, 2)
        })
    return sorted(archives, key=lambda x: x['filename'], reverse=True)

@app.get("/api/archives/{filename}")
def get_archive(filename: str):
    """Retrieve a specific archived transcript."""
    safe_path = os.path.join(ARCHIVE_DIR, os.path.basename(filename))
    if not os.path.exists(safe_path):
        raise HTTPException(status_code=404, detail="Archive not found")
    with open(safe_path, "r", encoding="utf-8") as f:
        return json.load(f)

@app.delete("/api/archives/{filename}")
def delete_archive(filename: str):
    """Delete a specific archive file."""
    safe_path = os.path.join(ARCHIVE_DIR, os.path.basename(filename))
    if os.path.exists(safe_path):
        os.remove(safe_path)
        return {"status": "deleted"}
    raise HTTPException(status_code=404, detail="File not found")

@app.post("/api/archives/prune")
def prune_archives(days: int = 180):
    """Delete files older than 'days'."""
    cutoff = time.time() - (days * 86400)
    deleted_count = 0
    files = glob.glob(os.path.join(ARCHIVE_DIR, "*.json"))
    
    for f in files:
        if os.stat(f).st_mtime < cutoff:
            os.remove(f)
            deleted_count += 1
            
    return {"status": "success", "deleted": deleted_count, "retention_days": days}

# --- EXPORT ENDPOINTS ---

@app.post("/api/download-pdf")
async def download_pdf(transcription_data: List[TranscriptionSegment], clean: bool = False):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    title = "Meeting Minutes Draft" if clean else "Audio Transcription"
    story.append(Paragraph(title, styles['h1']))
    story.append(Spacer(1, 0.2 * 100))

    if clean:
        # Paragraph Mode
        current_block = []
        for i, segment in enumerate(transcription_data):
            current_block.append(segment.text.strip())
            if (i + 1) % 5 == 0:
                text = " ".join(current_block)
                story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 0.1 * 100))
                current_block = []
        if current_block:
            story.append(Paragraph(" ".join(current_block), styles['Normal']))
    else:
        # Timestamp Mode
        for segment in transcription_data:
            text = f"<b>[{segment.start} - {segment.end}]:</b> {segment.text}"
            story.append(Paragraph(text, styles['Normal']))
            story.append(Spacer(1, 0.1 * 100))

    doc.build(story)
    buffer.seek(0)

    return Response(content=buffer.getvalue(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=\"transcription.pdf\""})


@app.post("/api/download-docx")
async def download_docx(transcription_data: List[TranscriptionSegment], clean: bool = False):
    document = Document()
    title = "Meeting Minutes Draft" if clean else "Audio Transcription"
    document.add_heading(title, level=1)

    if clean:
        current_block = []
        for i, segment in enumerate(transcription_data):
            current_block.append(segment.text.strip())
            if (i + 1) % 5 == 0:
                document.add_paragraph(" ".join(current_block))
                current_block = []
        if current_block:
            document.add_paragraph(" ".join(current_block))
    else:
        for segment in transcription_data:
            document.add_paragraph(f"[{segment.start} - {segment.end}] {segment.text}")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return Response(content=buffer.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=\"transcription.docx\""})

def download_youtube_audio(url: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    output_template = f"temp_yt_{timestamp}"
    
    ydl_opts = {
        'format': 'bestaudio/best',
        'outtmpl': output_template,
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'm4a',
        }],
        'quiet': True,
        'no_warnings': True,
        'source_address': '0.0.0.0',
    }
    
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=False)
        title = info.get('title', 'YouTube_Video')
        ydl.download([url])
        final_filename = f"{output_template}.m4a"
        return final_filename, title

@app.post("/api/transcribe")
async def transcribe_audio(
    file: UploadFile = File(None), 
    youtube_url: str = Form(None),
    model_size: str = Form("large-v2"),
    meeting_date: str = Form(None)
):
    if processing_lock.locked():
        raise HTTPException(
            status_code=503, 
            detail="System Busy: Another meeting is being processed. Please try again in ~20 minutes."
        )

    async with processing_lock:
        if not file and not youtube_url:
            raise HTTPException(status_code=400, detail="Please provide either a file or a YouTube URL.")

        print(f"Received request. Model: {model_size}, Date: {meeting_date}")
        
        temp_audio_path = None
        original_filename = "Unknown"

        try:
            if youtube_url:
                print(f"Downloading YouTube URL: {youtube_url}")
                temp_audio_path, video_title = download_youtube_audio(youtube_url)
                safe_title = "".join([c for c in video_title if c.isalnum() or c in (' ', '.', '_')]).replace(" ", "_")
                original_filename = f"{safe_title}.m4a"
            else:
                original_filename = file.filename
                temp_audio_path = f"temp_{file.filename}"
                with open(temp_audio_path, "wb") as buffer:
                    buffer.write(await file.read())

            if MOCK_MODE:
                await asyncio.sleep(2)
                result_segments = [
                    {"start": 0.0, "end": 5.0, "text": "This is a mock transcription."},
                    {"start": 5.0, "end": 10.0, "text": "Real GPU inference is disabled."}
                ]
            else:
                import whisperx
                device = "cuda" if torch.cuda.is_available() else "cpu"
                print(f"Loading WhisperX model {model_size} on {device}...")
                
                compute_type = "int8" 
                batch_size = 16
                
                model = whisperx.load_model(model_size, device, compute_type=compute_type)
                audio = whisperx.load_audio(temp_audio_path)
                result = model.transcribe(audio, batch_size=batch_size)
                
                model_lang = result["language"]
                del model
                gc.collect()
                if device == "cuda":
                    torch.cuda.empty_cache()

                print("Aligning...")
                model_a, metadata = whisperx.load_align_model(language_code=model_lang, device=device)
                result = whisperx.align(result["segments"], model_a, metadata, audio, device, return_char_alignments=False)
                
                del model_a
                gc.collect()
                if device == "cuda":
                    torch.cuda.empty_cache()
                
                result_segments = result["segments"]

            formatted_result = []
            for segment in result_segments:
                formatted_result.append({
                    "start": format_timestamp(segment.get("start", 0)),
                    "end": format_timestamp(segment.get("end", 0)),
                    "text": segment.get("text", "").strip()
                })

            try:
                date_prefix = meeting_date if meeting_date else datetime.now().strftime("%Y-%m-%d")
                time_suffix = datetime.now().strftime("%H-%M-%S")
                safe_filename = "".join([c for c in original_filename if c.isalnum() or c in (' ', '.', '_')]).replace(" ", "_")
                archive_path = os.path.join(ARCHIVE_DIR, f"{date_prefix}_{time_suffix}_{safe_filename}.json")
                
                with open(archive_path, "w", encoding="utf-8") as f:
                    json.dump(formatted_result, f, indent=2, ensure_ascii=False)
                print(f"Archived to: {archive_path}")
            except Exception as e:
                print(f"Archive failed: {e}")

            return {"transcription": formatted_result}

        except Exception as e:
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=str(e))
        
        finally:
            if temp_audio_path and os.path.exists(temp_audio_path):
                os.remove(temp_audio_path)

if os.path.exists("static"):
    app.mount("/static", StaticFiles(directory="static/static"), name="static")
    @app.get("/{full_path:path}")
    async def serve_react(full_path: str):
        if full_path.startswith("api"):
            return {"error": "API route not found"}
        return FileResponse("static/index.html")
else:
    print("Warning: 'static' directory not found. Frontend will not be served.")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
